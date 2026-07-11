"""Resume upload and listing.

Accepts PDF (recommended), DOCX, plain text, and resume images. The text
formats parse locally with no AI spend, so a bad file fails cheaply; images
are the one exception — a single vision call transcribes them to plain text
at upload, so everything downstream stays text-based. Extraction to a
structured profile happens lazily at session creation (see routes/sessions.py).
"""

import base64
from io import BytesIO
from pathlib import PurePosixPath

from docx import Document
from fastapi import APIRouter, Depends, HTTPException, UploadFile
from pypdf import PdfReader
from sqlmodel import Session, select

from app.agent.extraction import transcribe_resume_image
from app.db import get_session
from app.models import Resume

router = APIRouter(prefix="/api/resumes", tags=["resumes"])

IMAGE_MEDIA_TYPES = {
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".webp": "image/webp",
}
ACCEPTED = "a PDF (recommended), DOCX, TXT, or an image (PNG, JPG, WebP)"


def _extract_pdf_text(data: bytes) -> str:
    reader = PdfReader(BytesIO(data))
    return "\n".join(page.extract_text() or "" for page in reader.pages).strip()


def _extract_docx_text(data: bytes) -> str:
    doc = Document(BytesIO(data))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(part for part in parts if part.strip()).strip()


@router.post("")
async def upload_resume(
    file: UploadFile, session: Session = Depends(get_session)
) -> dict:
    suffix = PurePosixPath((file.filename or "").lower()).suffix
    data = await file.read()

    if suffix in IMAGE_MEDIA_TYPES:
        # The one upload path that costs an AI call — errors here are the
        # gateway's to surface, not a "bad file" 400.
        image_b64 = base64.standard_b64encode(data).decode("ascii")
        text = transcribe_resume_image(image_b64, IMAGE_MEDIA_TYPES[suffix]).strip()
    elif suffix == ".pdf":
        try:
            text = _extract_pdf_text(data)
        except Exception:
            raise HTTPException(status_code=400, detail="Could not read the PDF file.")
    elif suffix == ".docx":
        try:
            text = _extract_docx_text(data)
        except Exception:
            raise HTTPException(status_code=400, detail="Could not read the Word document.")
    elif suffix in (".txt", ".md"):
        text = data.decode("utf-8", errors="replace").strip()
    else:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type. Upload {ACCEPTED}.",
        )

    if not text:
        raise HTTPException(
            status_code=400,
            detail=(
                "No text could be extracted from the file. If it's a scanned "
                "PDF, upload a screenshot of it instead."
            ),
        )

    resume = Resume(filename=file.filename, raw_text=text)
    session.add(resume)
    session.commit()
    session.refresh(resume)
    return {"id": resume.id, "filename": resume.filename}


@router.get("")
def list_resumes(session: Session = Depends(get_session)) -> list[dict]:
    resumes = session.exec(select(Resume)).all()
    return [{"id": r.id, "filename": r.filename} for r in resumes]
